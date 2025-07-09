#!/usr/bin/env python3

import markdown
from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.oxml.shared import OxmlElement, qn
import html.parser

# Test HTML content with code blocks
test_html = """<h1>Test Document</h1>
<p>Inline code: <code>print("hello")</code></p>
<p>Code block:</p>
<div class="codehilite"><pre><span></span><code><span class="k">def</span> <span class="nf">hello</span><span class="p">():</span>
    <span class="nb">print</span><span class="p">(</span><span class="s2">&quot;Hello, World!&quot;</span><span class="p">)</span>
    <span class="k">return</span> <span class="kc">True</span>
</code></pre></div>
<p>Another paragraph.</p>"""

class DebugCodeParser(html.parser.HTMLParser):
    def __init__(self, document):
        super().__init__()
        self.doc = document
        self.current_paragraph = None
        self.in_code = False
        self.in_heading = False
        self.heading_level = 1
        self.text_buffer = ""
        self.in_code_block = False
        
    def handle_starttag(self, tag, attrs):
        print(f"START TAG: {tag} (attrs: {attrs})")
        if tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            self.flush_text()
            self.in_heading = True
            self.heading_level = int(tag[1])
        elif tag == 'p':
            self.flush_text()
            self.current_paragraph = self.doc.add_paragraph()
        elif tag == 'code':
            print("  -> ENTERING CODE MODE")
            # Flush any existing text before entering code mode
            self.flush_text()
            self.in_code = True
        elif tag == 'div':
            # Check if this is a code block div
            for attr_name, attr_value in attrs:
                if attr_name == 'class' and 'codehilite' in attr_value:
                    print("  -> ENTERING CODE BLOCK MODE")
                    self.flush_text()
                    self.in_code_block = True
                    self.current_paragraph = self.doc.add_paragraph()
                    break
        elif tag == 'pre':
            if self.in_code_block:
                print("  -> PRE inside code block - setting in_code = True")
                self.in_code = True
            else:
                self.flush_text()
                self.current_paragraph = self.doc.add_paragraph()
        elif tag == 'span':
            print("  -> SPAN (ignoring for syntax highlighting)")
            pass
            
    def handle_endtag(self, tag):
        print(f"END TAG: {tag}")
        if tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            self.flush_text_as_heading()
            self.in_heading = False
        elif tag == 'p':
            self.flush_text()
        elif tag == 'code':
            # Always flush text when closing code tag (for both inline and block code)
            print("  -> CODE end - flushing")
            self.flush_text()
            print("  -> EXITING CODE MODE")
            self.in_code = False
        elif tag == 'pre':
            if self.in_code_block:
                print("  -> PRE end inside code block")
                self.flush_text()
                self.in_code = False
            else:
                self.flush_text()
        elif tag == 'div':
            if self.in_code_block:
                print("  -> EXITING CODE BLOCK MODE")
                self.in_code_block = False
                self.current_paragraph = None
        elif tag == 'span':
            pass  # Ignore span end tags
            
    def handle_data(self, data):
        print(f"DATA: '{repr(data)}'")
        self.text_buffer += data
        
    def flush_text(self):
        if self.text_buffer.strip():
            print(f"FLUSHING TEXT (in_code={self.in_code}): '{self.text_buffer.strip()}'")
            if not self.current_paragraph:
                self.current_paragraph = self.doc.add_paragraph()
            
            if self.in_code:
                run = self.current_paragraph.add_run(self.text_buffer)
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
                # Set white text color
                run.font.color.rgb = RGBColor(255, 255, 255)
                # Set grey background using shading
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:val'), 'clear')
                shading_elm.set(qn('w:color'), 'auto')
                shading_elm.set(qn('w:fill'), '808080')
                run._element.rPr.append(shading_elm)
                print("  -> Applied code formatting!")
            else:
                self.current_paragraph.add_run(self.text_buffer)
        self.text_buffer = ""
        
    def flush_text_as_heading(self):
        if self.text_buffer.strip():
            print(f"FLUSHING HEADING: '{self.text_buffer.strip()}'")
            heading = self.doc.add_heading(self.text_buffer.strip(), level=self.heading_level)
        self.text_buffer = ""
        self.current_paragraph = None

# Test the parser
doc = Document()
parser = DebugCodeParser(doc)
parser.feed(test_html)
parser.flush_text()

# Save the document
doc.save('debug_code_blocks_output.docx')
print("\nDocument saved as debug_code_blocks_output.docx")

# Check the results
print(f"\nDocument has {len(doc.paragraphs)} paragraphs:")
for i, para in enumerate(doc.paragraphs):
    print(f"  Paragraph {i}: '{para.text}'")
    for j, run in enumerate(para.runs):
        is_code = run.font.name == 'Courier New'
        print(f"    Run {j}: '{run.text}' (code={is_code})")
