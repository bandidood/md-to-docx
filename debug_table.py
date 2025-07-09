#!/usr/bin/env python3

import markdown
from docx import Document
from docx.shared import Inches
import html.parser

# Test HTML content
test_html = """<h1>Test Document</h1>
<p>Simple paragraph before table.</p>
<table>
<thead>
<tr>
<th>Name</th>
<th>Age</th>
<th>City</th>
</tr>
</thead>
<tbody>
<tr>
<td>John</td>
<td>25</td>
<td>Paris</td>
</tr>
<tr>
<td>Jane</td>
<td>30</td>
<td>London</td>
</tr>
</tbody>
</table>
<p>Simple paragraph after table.</p>"""

class DebugHTMLToDocxParser(html.parser.HTMLParser):
    def __init__(self, document):
        super().__init__()
        self.doc = document
        self.current_paragraph = None
        self.in_code = False
        self.in_heading = False
        self.heading_level = 1
        self.list_level = 0
        self.text_buffer = ""
        # Table handling
        self.current_table = None
        self.current_row = None
        self.current_cell = None
        self.in_table_header = False
        
    def handle_starttag(self, tag, attrs):
        print(f"START TAG: {tag}")
        if tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            self.flush_text()
            self.in_heading = True
            self.heading_level = int(tag[1])
        elif tag == 'p':
            self.flush_text()
            self.current_paragraph = self.doc.add_paragraph()
        elif tag == 'table':
            print("  CREATING TABLE")
            self.flush_text()
            self.current_table = self.doc.add_table(rows=0, cols=0)
            self.current_table.style = 'Table Grid'
            print(f"  Table created: {self.current_table}")
        elif tag == 'thead':
            self.in_table_header = True
        elif tag == 'tbody':
            self.in_table_header = False
        elif tag == 'tr':
            if self.current_table is not None:
                print("  ADDING ROW")
                self.current_row = self.current_table.add_row()
                self.current_cell_index = 0
                print(f"  Row added: {self.current_row}")
        elif tag in ['td', 'th']:
            if self.current_row is not None:
                print(f"  ADDING CELL {self.current_cell_index}")
                # Ensure table has enough columns
                while len(self.current_table.columns) <= self.current_cell_index:
                    print(f"    Adding column {len(self.current_table.columns)}")
                    self.current_table.add_column(Inches(1.5))
                
                # Get the current cell
                if self.current_cell_index < len(self.current_row.cells):
                    self.current_cell = self.current_row.cells[self.current_cell_index]
                    print(f"    Cell selected: {self.current_cell}")
                    # Clear the cell and prepare for content
                    self.current_cell.text = ''
                    self.current_paragraph = self.current_cell.paragraphs[0]
                    
    def handle_endtag(self, tag):
        print(f"END TAG: {tag}")
        if tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            self.flush_text_as_heading()
            self.in_heading = False
        elif tag == 'p':
            self.flush_text()
        elif tag == 'table':
            print("  TABLE ENDED")
            self.current_table = None
            self.current_row = None
            self.current_cell = None
        elif tag == 'tr':
            print("  ROW ENDED")
            self.current_row = None
        elif tag in ['td', 'th']:
            print(f"  CELL ENDED, incrementing index from {getattr(self, 'current_cell_index', 'None')}")
            # Flush any remaining text in the cell before closing
            self.flush_text()
            if hasattr(self, 'current_cell_index'):
                self.current_cell_index += 1
            self.current_cell = None
            self.current_paragraph = None
            
    def handle_data(self, data):
        print(f"DATA: '{data.strip()}'")
        self.text_buffer += data
        
    def flush_text(self):
        if self.text_buffer.strip():
            print(f"FLUSHING TEXT: '{self.text_buffer.strip()}'")
            if not self.current_paragraph:
                self.current_paragraph = self.doc.add_paragraph()
            self.current_paragraph.add_run(self.text_buffer)
        self.text_buffer = ""
        
    def flush_text_as_heading(self):
        if self.text_buffer.strip():
            print(f"FLUSHING HEADING: '{self.text_buffer.strip()}'")
            heading = self.doc.add_heading(self.text_buffer.strip(), level=self.heading_level)
        self.text_buffer = ""
        self.current_paragraph = None

# Test the parser
doc = Document()
parser = DebugHTMLToDocxParser(doc)
parser.feed(test_html)
parser.flush_text()

# Save the document
doc.save('debug_table_output.docx')
print("\nDocument saved as debug_table_output.docx")
print(f"Number of tables in document: {len(doc.tables)}")
if doc.tables:
    table = doc.tables[0]
    print(f"Table rows: {len(table.rows)}")
    print(f"Table columns: {len(table.columns)}")
    for i, row in enumerate(table.rows):
        print(f"Row {i}: {[cell.text for cell in row.cells]}")
