#!/usr/bin/env python3
"""
Markdown to DOCX Converter with Mermaid Support

This script converts Markdown files to DOCX format with embedded Mermaid diagrams.
Mermaid diagrams are converted to images and embedded in the DOCX file.

Usage:
    python md_to_docx.py input.md output.docx

Requirements:
    - python-docx
    - markdown
    - requests
    - pillow
    - Node.js with @mermaid-js/mermaid-cli (mmdc)
"""

import argparse
import os
import re
import sys
import tempfile
import subprocess
from pathlib import Path
import markdown
from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.oxml.shared import OxmlElement, qn
from PIL import Image
import html


def find_mmdc():
    """Find mmdc executable on the system"""
    # Common locations to check for mmdc
    possible_paths = [
        'mmdc',  # If in PATH
        'mmdc.cmd',  # Windows
    ]
    
    # Check npm global directory on Windows
    try:
        npm_prefix = subprocess.run(['npm', 'config', 'get', 'prefix'], 
                                  capture_output=True, text=True, check=True)
        npm_bin = os.path.join(npm_prefix.stdout.strip(), 'mmdc.cmd')
        possible_paths.append(npm_bin)
    except:
        pass
    
    # Also check user AppData on Windows
    if os.name == 'nt':  # Windows
        user_npm = os.path.expanduser('~\\AppData\\Roaming\\npm\\mmdc.cmd')
        possible_paths.append(user_npm)
    
    for path in possible_paths:
        try:
            result = subprocess.run([path, '--version'], 
                                  capture_output=True, text=True)
            if result.returncode == 0:
                return path
        except (subprocess.CalledProcessError, FileNotFoundError):
            continue
    
    return None


class MarkdownToDocxConverter:
    def __init__(self):
        self.temp_files = []
        self.mmdc_cmd = find_mmdc()
        
    def __del__(self):
        """Cleanup temporary files"""
        self.cleanup_temp_files()
        
    def cleanup_temp_files(self):
        """Remove temporary files"""
        for temp_file in self.temp_files:
            try:
                if os.path.exists(temp_file):
                    os.remove(temp_file)
            except Exception as e:
                print(f"Warning: Could not remove temp file {temp_file}: {e}")
        self.temp_files.clear()
        
    def extract_mermaid_diagrams(self, markdown_content):
        """Extract Mermaid diagrams from markdown content"""
        # Pattern to match mermaid code blocks
        mermaid_pattern = r'```mermaid\n(.*?)\n```'
        
        diagrams = []
        matches = re.finditer(mermaid_pattern, markdown_content, re.DOTALL)
        
        for i, match in enumerate(matches):
            diagram_code = match.group(1).strip()
            diagrams.append({
                'index': i,
                'code': diagram_code,
                'full_match': match.group(0),
                'start': match.start(),
                'end': match.end()
            })
            
        return diagrams
        
    def generate_mermaid_image(self, mermaid_code, output_path):
        """Generate PNG image from Mermaid code using mmdc CLI"""
        try:
            # Create temporary file for mermaid code with UTF-8 encoding
            with tempfile.NamedTemporaryFile(mode='w', suffix='.mmd', delete=False, encoding='utf-8') as temp_mmd:
                temp_mmd.write(mermaid_code)
                temp_mmd_path = temp_mmd.name
                self.temp_files.append(temp_mmd_path)
            
            # Run mmdc to generate PNG
            cmd = [
                self.mmdc_cmd,
                '-i', temp_mmd_path,
                '-o', output_path,
                '-t', 'default',  # default theme (colored)
                '-b', 'white',    # white background
                '--width', '1200',
                '--height', '800'
            ]
            
            result = subprocess.run(cmd, capture_output=True, text=True, encoding='utf-8')
            
            if result.returncode != 0:
                print(f"Error generating Mermaid diagram: {result.stderr}")
                if result.stdout:
                    print(f"Stdout: {result.stdout}")
                return False
                
            return os.path.exists(output_path)
            
        except Exception as e:
            print(f"Error generating Mermaid image: {e}")
            return False
            
    def process_mermaid_diagrams(self, markdown_content):
        """Process and replace Mermaid diagrams with placeholders"""
        diagrams = self.extract_mermaid_diagrams(markdown_content)
        
        if not diagrams:
            return markdown_content, []
            
        # Generate images for each diagram
        diagram_images = []
        processed_content = markdown_content
        
        # Process diagrams in reverse order to maintain correct indices
        for diagram in reversed(diagrams):
            # Create temporary image file
            image_path = tempfile.mktemp(suffix='.png')
            self.temp_files.append(image_path)
            
            if self.generate_mermaid_image(diagram['code'], image_path):
                diagram_images.insert(0, {
                    'path': image_path,
                    'index': diagram['index']
                })
                
                # Replace mermaid code block with placeholder
                placeholder = f"[MERMAID_DIAGRAM_{diagram['index']}]"
                processed_content = (processed_content[:diagram['start']] + 
                                   placeholder + 
                                   processed_content[diagram['end']:])
            else:
                print(f"Failed to generate image for diagram {diagram['index']}")
                
        return processed_content, diagram_images
        
    def markdown_to_docx_elements(self, markdown_content, doc):
        """Convert markdown content to DOCX elements"""
        # Convert markdown to HTML first
        md = markdown.Markdown(extensions=['tables', 'fenced_code', 'codehilite'])
        html_content = md.convert(markdown_content)
        
        # Parse HTML and convert to DOCX
        self.html_to_docx(html_content, doc)
        
    def html_to_docx(self, html_content, doc):
        """Convert HTML content to DOCX elements"""
        # Simple HTML to DOCX conversion
        # This is a basic implementation - you might want to use a more sophisticated library
        
        # Remove HTML tags and convert to plain text with basic formatting
        import html.parser
        
        class HTMLToDocxParser(html.parser.HTMLParser):
            def __init__(self, document):
                super().__init__()
                self.doc = document
                self.current_paragraph = None
                self.in_code = False
                self.in_heading = False
                self.heading_level = 1
                self.list_level = 0
                self.text_buffer = ""
                # Table handling
                self.current_table = None
                self.current_row = None
                self.current_cell = None
                self.in_table_header = False
                # Code block handling
                self.in_code_block = False
                self.current_language = None
                
            def handle_starttag(self, tag, attrs):
                if tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                    self.flush_text()
                    self.in_heading = True
                    self.heading_level = int(tag[1])
                elif tag == 'p':
                    self.flush_text()
                    self.current_paragraph = self.doc.add_paragraph()
                elif tag == 'code':
                    # Flush any existing text before entering code mode
                    self.flush_text()
                    self.in_code = True
                elif tag == 'div':
                    # Check if this is a code block div
                    for attr_name, attr_value in attrs:
                        if attr_name == 'class' and 'codehilite' in attr_value:
                            self.flush_text()
                            self.in_code_block = True
                            self.current_paragraph = self.doc.add_paragraph()
                            break
                elif tag == 'pre':
                    if self.in_code_block:
                        # Inside a code block, treat everything as code
                        self.in_code = True
                        # Ajouter un peu d'espacement avant le bloc de code
                        if self.current_paragraph:
                            self.current_paragraph.space_before = Pt(6)
                            self.current_paragraph.space_after = Pt(6)
                    else:
                        self.flush_text()
                        self.current_paragraph = self.doc.add_paragraph()
                elif tag in ['ul', 'ol']:
                    self.flush_text()
                    self.list_level += 1
                elif tag == 'li':
                    self.flush_text()
                    self.current_paragraph = self.doc.add_paragraph(style='List Bullet')
                elif tag == 'table':
                    self.flush_text()
                    self.current_table = self.doc.add_table(rows=0, cols=0)
                    self.current_table.style = 'Table Grid'
                elif tag == 'thead':
                    self.in_table_header = True
                elif tag == 'tbody':
                    self.in_table_header = False
                elif tag == 'tr':
                    if self.current_table is not None:
                        self.current_row = self.current_table.add_row()
                        self.current_cell_index = 0  # Reset cell index for new row
                elif tag in ['td', 'th']:
                    if self.current_row is not None:
                        # Determine current cell index
                        self.current_cell_index = getattr(self, 'current_cell_index', 0)
                        
                        # Ensure table has enough columns
                        while len(self.current_table.columns) <= self.current_cell_index:
                            self.current_table.add_column(Inches(1.5))
                        
                        # Get the current cell
                        if self.current_cell_index < len(self.current_row.cells):
                            self.current_cell = self.current_row.cells[self.current_cell_index]
                            # Clear the cell and prepare for content
                            self.current_cell.text = ''
                            self.current_paragraph = self.current_cell.paragraphs[0]
                elif tag == 'br':
                    self.text_buffer += '\n'
                elif tag == 'span':
                    # Ignore span tags (used for syntax highlighting) but keep the content
                    pass
                    
            def handle_endtag(self, tag):
                if tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                    self.flush_text_as_heading()
                    self.in_heading = False
                elif tag == 'p':
                    self.flush_text()
                elif tag == 'code':
                    # Always flush text when closing code tag (for both inline and block code)
                    self.flush_text()
                    self.in_code = False
                elif tag == 'pre':
                    if self.in_code_block:
                        # End of code content inside code block
                        self.flush_text()
                        self.in_code = False
                    else:
                        self.flush_text()
                elif tag == 'div':
                    if self.in_code_block:
                        # End of code block
                        self.in_code_block = False
                        self.current_paragraph = None
                elif tag in ['ul', 'ol']:
                    self.list_level -= 1
                elif tag == 'li':
                    self.flush_text()
                elif tag == 'table':
                    self.current_table = None
                    self.current_row = None
                    self.current_cell = None
                elif tag == 'tr':
                    self.current_row = None
                elif tag in ['td', 'th']:
                    # Flush any remaining text in the cell before closing
                    self.flush_text()
                    if hasattr(self, 'current_cell_index'):
                        self.current_cell_index += 1
                    self.current_cell = None
                    self.current_paragraph = None
                    
            def handle_data(self, data):
                self.text_buffer += data
                
            def flush_text(self):
                if self.text_buffer.strip():
                    if not self.current_paragraph:
                        self.current_paragraph = self.doc.add_paragraph()
                    
                    if self.in_code:
                        run = self.current_paragraph.add_run(self.text_buffer)
                        run.font.name = 'Courier New'  # Police Courier New pour les blocs de code
                        run.font.size = Pt(9)  # Taille de police légèrement plus petite
                        
                        # Pour les blocs de code (pas les codes Mermaid)
                        if self.in_code_block:
                            # Texte sombre pour les blocs de code
                            run.font.color.rgb = RGBColor(33, 37, 41)  # Texte sombre (#212529)
                            
                            # Appliquer le fond gris clair à tout le paragraphe
                            p_pr = self.current_paragraph._element.get_or_add_pPr()
                            
                            # Fond gris clair pour tout le paragraphe
                            shading_elm = OxmlElement('w:shd')
                            shading_elm.set(qn('w:val'), 'clear')
                            shading_elm.set(qn('w:color'), 'auto')
                            shading_elm.set(qn('w:fill'), 'E9ECEF')  # Fond gris clair (#E9ECEF)
                            p_pr.append(shading_elm)
                            
                            # Ajouter une bordure au paragraphe pour délimiter le bloc
                            p_borders = OxmlElement('w:pBdr')
                            
                            # Bordure gauche plus épaisse pour l'effet "bloc de code"
                            left_border = OxmlElement('w:left')
                            left_border.set(qn('w:val'), 'single')
                            left_border.set(qn('w:sz'), '12')  # Épaisseur
                            left_border.set(qn('w:space'), '4')
                            left_border.set(qn('w:color'), '569CD6')  # Couleur bleue
                            p_borders.append(left_border)
                            p_pr.append(p_borders)
                            
                            # Indentation pour le bloc de code
                            ind = OxmlElement('w:ind')
                            ind.set(qn('w:left'), '340')  # Indentation gauche
                            p_pr.append(ind)
                            
                        else:
                            # Code inline - style plus discret
                            run.font.color.rgb = RGBColor(86, 156, 214)  # Bleu pour code inline
                            shading_elm = OxmlElement('w:shd')
                            shading_elm.set(qn('w:val'), 'clear')
                            shading_elm.set(qn('w:color'), 'auto')
                            shading_elm.set(qn('w:fill'), 'F5F5F5')  # Fond gris très clair
                            run._element.rPr.append(shading_elm)
                    else:
                        self.current_paragraph.add_run(self.text_buffer)
                        
                self.text_buffer = ""
                
            def flush_text_as_heading(self):
                if self.text_buffer.strip():
                    heading = self.doc.add_heading(self.text_buffer.strip(), level=self.heading_level)
                self.text_buffer = ""
                self.current_paragraph = None
        
        parser = HTMLToDocxParser(doc)
        parser.feed(html_content)
        parser.flush_text()
        
    def insert_mermaid_images(self, doc, diagram_images):
        """Insert Mermaid diagram images into the document"""
        # Find placeholders and replace with images
        for paragraph in doc.paragraphs:
            for diagram in diagram_images:
                placeholder = f"[MERMAID_DIAGRAM_{diagram['index']}]"
                if placeholder in paragraph.text:
                    # Clear the paragraph text
                    paragraph.clear()
                    
                    # Add the image
                    try:
                        # Check if image exists and get dimensions
                        if os.path.exists(diagram['path']):
                            with Image.open(diagram['path']) as img:
                                width, height = img.size
                                
                            # Calculate appropriate size for document
                            max_width = Inches(6)  # Max width in document
                            aspect_ratio = height / width
                            
                            if width > max_width.inches * 96:  # 96 DPI
                                doc_width = max_width
                                doc_height = Inches(max_width.inches * aspect_ratio)
                            else:
                                doc_width = Inches(width / 96)
                                doc_height = Inches(height / 96)
                            
                            run = paragraph.add_run()
                            run.add_picture(diagram['path'], width=doc_width, height=doc_height)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            
                    except Exception as e:
                        print(f"Error inserting image for diagram {diagram['index']}: {e}")
                        paragraph.add_run(f"[Error loading Mermaid diagram {diagram['index']}]")
                        
    def convert(self, input_file, output_file):
        """Convert Markdown file to DOCX with Mermaid diagrams"""
        try:
            # Read input file
            with open(input_file, 'r', encoding='utf-8') as f:
                markdown_content = f.read()
            
            print(f"Processing {input_file}...")
            
            # Process Mermaid diagrams
            processed_content, diagram_images = self.process_mermaid_diagrams(markdown_content)
            
            if diagram_images:
                print(f"Generated {len(diagram_images)} Mermaid diagrams")
            
            # Create new document
            doc = Document()
            
            # Convert markdown to DOCX
            self.markdown_to_docx_elements(processed_content, doc)
            
            # Insert Mermaid images
            if diagram_images:
                self.insert_mermaid_images(doc, diagram_images)
            
            # Save document
            doc.save(output_file)
            print(f"Converted successfully: {output_file}")
            
            return True
            
        except Exception as e:
            print(f"Error converting file: {e}")
            return False
        finally:
            self.cleanup_temp_files()


def main():
    parser = argparse.ArgumentParser(description='Convert Markdown to DOCX with Mermaid diagram support')
    parser.add_argument('input', help='Input Markdown file')
    parser.add_argument('output', help='Output DOCX file')
    
    args = parser.parse_args()
    
    # Check if input file exists
    if not os.path.exists(args.input):
        print(f"Error: Input file '{args.input}' not found")
        sys.exit(1)
    
    # Check if mmdc is available
    mmdc_cmd = find_mmdc()
    if not mmdc_cmd:
        print("Error: Mermaid CLI (mmdc) not found. Please install it with:")
        print("npm install -g @mermaid-js/mermaid-cli")
        sys.exit(1)
    
    # Convert file
    converter = MarkdownToDocxConverter()
    success = converter.convert(args.input, args.output)
    
    if not success:
        sys.exit(1)


if __name__ == "__main__":
    main()